import hashlib
from pathlib import Path
from typing import List

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from app.core.config import settings
from app.core.logger import get_logger
from app.core.vector_store import VectorStoreManager

logger = get_logger(__name__)


def _make_doc_id(content: str, metadata: dict) -> str:
    payload = content + str(sorted(metadata.items()))
    return hashlib.sha256(payload.encode()).hexdigest()[:16]


def _load_pdf(path: str) -> str:
    try:
        from pypdf import PdfReader
        reader = PdfReader(path)
        text = "\n\n".join([page.extract_text() or "" for page in reader.pages])
        logger.info("pdf_loaded", pages=len(reader.pages))
        return text
    except ImportError:
        raise ImportError("Run: pip install pypdf")


def _load_docx(path: str) -> str:
    try:
        import docx
        doc = docx.Document(path)
        text = "\n\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        logger.info("docx_loaded", paragraphs=len(doc.paragraphs))
        return text
    except ImportError:
        raise ImportError("Run: pip install python-docx")


def _load_text(path: str) -> str:
    return Path(path).read_text(encoding="utf-8")


class IngestionService:
    def __init__(self) -> None:
        self.splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.CHUNK_SIZE,
            chunk_overlap=settings.CHUNK_OVERLAP,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""],
        )
        self.vs_manager = VectorStoreManager()

    def ingest_texts(self, texts: List[str], metadatas: List[dict] | None = None) -> dict:
        if metadatas is None:
            metadatas = [{} for _ in texts]
        if len(texts) != len(metadatas):
            raise ValueError("texts and metadatas must have the same length.")

        raw_docs = [Document(page_content=t, metadata=m) for t, m in zip(texts, metadatas)]
        chunks = self.splitter.split_documents(raw_docs)
        logger.info("text_split", raw_docs=len(raw_docs), chunks=len(chunks))
        stored_ids = self.vs_manager.add_documents(chunks)
        return {"raw_document_count": len(raw_docs), "chunk_count": len(chunks), "stored_ids": stored_ids}

    def ingest_file(self, file_path: str, original_filename: str | None = None, ext: str | None = None, extra_metadata: dict | None = None) -> dict:
        path = Path(file_path)
        ext = ext or path.suffix.lower()
        name = original_filename or path.name

        loaders = {".pdf": _load_pdf, ".docx": _load_docx, ".txt": _load_text, ".md": _load_text}
        if ext not in loaders:
            raise ValueError(f"Unsupported file type: {ext}")

        text = loaders[ext](str(path))
        if not text.strip():
            raise ValueError(f"File '{name}' is empty or unreadable.")

        metadata = {"source": name, "filename": name, "file_type": ext.lstrip(".")}
        if extra_metadata:
            metadata.update(extra_metadata)

        logger.info("file_loaded", filename=name, type=ext, chars=len(text))
        return self.ingest_texts([text], [metadata])